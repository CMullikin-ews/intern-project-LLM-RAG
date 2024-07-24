import subprocess
import docx
import streamlit as st
from langchain_community.document_loaders import DirectoryLoader, JSONLoader, TextLoader, UnstructuredFileLoader, UnstructuredHTMLLoader, UnstructuredMarkdownLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.llms import Ollama
from langchain.schema import Document
import os
import PyPDF2
import logging
import random
import re
import nltk
from collections import Counter
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Download NLTK stopwords
nltk.download('stopwords')
nltk.download('punkt')

# Define data paths
DATA_PATH = '../../cyber_data'
DB_FAISS_PATH = '../vectorstore'

# File loaders
loaders = {
    '.php': UnstructuredFileLoader,
    '.cs': UnstructuredFileLoader,
    '': UnstructuredFileLoader,
    '.c': UnstructuredFileLoader,
    '.html': UnstructuredHTMLLoader,
    '.md': UnstructuredMarkdownLoader,
    '.txt': TextLoader,
    '.ps1': UnstructuredFileLoader,
    '.delphi': UnstructuredFileLoader,
    '.asm': UnstructuredFileLoader,
    '.TXT': TextLoader
}

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def setup_ollama():
    try:
        subprocess.run("curl -fsSL https://ollama.com/install.sh | sh", shell=True, check=True)
        os.environ["OLLAMA_HOST"] = "localhost:8888"
        subprocess.run("sudo service ollama stop", shell=True, check=True)
        cmd = "ollama serve"
        process = subprocess.Popen(cmd, shell=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        logging.info("Ollama setup completed successfully.")
    except subprocess.CalledProcessError as e:
        logging.error(f"Error setting up Ollama: {e.stderr.decode()}")
    except Exception as e:
        logging.error(f"Unexpected error: {e}")

def get_file_types(directory):
    file_types = set()
    for filename in os.listdir(directory):
        if os.path.isfile(os.path.join(directory, filename)):
            _, ext = os.path.splitext(filename)
            file_types.add(ext.lower())
    return file_types

def create_directory_loader(file_type, directory_path):
    """
    Creates and returns a DirectoryLoader using the loader specific to the file type provided
    
    Args:
        file_type (str): Type of file to make loader for
        directory_path (str): Path to directory

    Returns:
        DirectoryLoader: loader for the files in the directory provided
    """
    if file_type == '.json':
        loader_list = []
        for file_name in [file for file in os.listdir(directory_path) if file.endswith('.json')]:
            loader_list.append(JSONLoader(file_path=os.path.join(directory_path, file_name), jq_schema='.', text_content=False))
        return loader_list
    else:
        return DirectoryLoader(
            path=directory_path,
            glob=f"**/*{file_type}",
            loader_cls=loaders.get(file_type, UnstructuredFileLoader))

def load_documents():
    file_types = get_file_types(DATA_PATH)
    documents = []
    
    for file_type in file_types:
        if file_type.strip() != "":
            logger.info(f"Processing files of type: {file_type}")
            try:
                if file_type == '.json':
                    loader_list = create_directory_loader(file_type, DATA_PATH)
                    for loader in loader_list:
                        docs = loader.load()
                        chunks = split_text(docs)
                        if chunks:
                            documents.extend(chunks)
                else:
                    loader = create_directory_loader(file_type, DATA_PATH)
                    docs = loader.load()
                    chunks = split_text(docs)
                    if chunks:
                        documents.extend(chunks)
            except Exception as e:
                logger.error(f"Error loading documents of type {file_type}: {e}")
    return documents

def split_text(docs, max_length=512, chunk_overlap=50):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=max_length,
        chunk_overlap=chunk_overlap
    )
    return splitter.split_documents(docs)

def create_knowledge_base():
    docs = load_documents()
    os.system("ollama pull mxbai-embed-large")
    chunks = split_text(docs)
    embeddings = OllamaEmbeddings(model="mxbai-embed-large", show_progress=True)
    documents = [Document(page_content=chunk) for chunk in chunks]
    vectorstore = FAISS.from_documents(documents=documents, embedding=embeddings, allow_dangerous_deserialization=True)
    vectorstore.save_local(DB_FAISS_PATH)

def load_knowledge_base():
    embeddings = OllamaEmbeddings(model="mxbai-embed-large", show_progress=True)
    db = FAISS.load_local(DB_FAISS_PATH, embeddings, allow_dangerous_deserialization=True)
    return db

def load_llm():
    try:
        return Ollama(model="llama3")
    except Exception as e:
        logger.error(f"Error loading LLM: {e}")
        return None

def load_prompt():
    prompt = """
    You are an assistant for helping software developers to detect and neutralize viruses.
    Make sure to clearly define any necessary terms and go through the steps to use any application or software.
    Only use the data provided to you.
    Cite the sources used in constructing the response.
    If the answer is not in the data provided, answer "Sorry, I'm not sure how to respond to this"
    """
    try:
        return ChatPromptTemplate.from_template(prompt)
    except Exception as e:
        logger.error(f"Error loading prompt: {e}")
        return None

def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

def generate_response(query: str) -> str:
    responses = [
        "Sure, I can help with that!",
        "Let me find that information for you.",
        "Here is what I found.",
        "This is the information you requested."
    ]
    return random.choice(responses)

def get_relevant_url(query: str) -> str:
    urls = [
        "https://example.com/info1",
        "https://example.com/info2",
        "https://example.com/info3",
        "https://example.com/info4"
    ]
    return random.choice(urls)

def respond_with_sources(query, retriever) -> str:
    try:
        retrieved_docs = retriever.invoke(query)
        sources = [doc.metadata.get('source', 'Unknown') for doc in retrieved_docs]
        citation_text = "Sources: " + ", ".join(sources)
        return f"\n\n{citation_text}"
    except Exception as e:
        logger.error(f"Error retrieving sources: {e}")
        return "\n\nSources: Unknown"

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ''
    
    try:
        if ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        
        elif ext == '.pdf':
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfFileReader(file)
                for page_num in range(reader.numPages):
                    text += reader.getPage(page_num).extract_text()
        
        elif ext == '.docx':
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'
    
    except Exception as e:
        logger.error(f"Error extracting text from file {file_path}: {e}")
    
    return text

def extract_text_from_directory(directory_path):
    extracted_texts = {}
    
    try:
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            if os.path.isfile(file_path):
                text = extract_text_from_file(file_path)
                extracted_texts[filename] = text
    except Exception as e:
        logger.error(f"Error extracting text from directory {directory_path}: {e}")
    
    return extracted_texts

def extract_metadata(text):
    """Extract metadata from the text."""
    metadata = {}
    
    # Extract title (assuming the title is the first line)
    lines = text.split('\n')
    metadata['title'] = lines[0] if lines else 'Unknown Title'
    
    # Extract author (assuming author is mentioned in the second line)
    metadata['author'] = lines[1] if len(lines) > 1 else 'Unknown Author'
    
    # Extract date (assuming date is mentioned in the third line in a known format)
    date_line = lines[2] if len(lines) > 2 else ''
    date_match = re.search(r'\b\d{4}-\d{2}-\d{2}\b', date_line)
    metadata['date'] = date_match.group(0) if date_match else 'Unknown Date'
    
    # Extract keywords (most common non-stopwords)
    stop_words = set(stopwords.words('english'))
    words = word_tokenize(text)
    words = [word.lower() for word in words if word.isalnum() and word.lower() not in stop_words]
    keywords = [word for word, freq in Counter(words).most_common(10)]
    metadata['keywords'] = keywords
    
    return metadata

if __name__ == '__main__':
    setup_ollama()
    st.header("Welcome to the 📝Computer Virus Copilot")
    st.write("🤖 You can chat by entering your queries")

    try:
        knowledge_base = load_knowledge_base()
        llm = load_llm()
        prompt = load_prompt()
        logging.info("Components loaded successfully.")
    except Exception as e:
        logging.error(f"Error loading components: {e}")
        st.write("An error occurred while loading the components. Please check the logs.")

    query = st.text_input('Enter some text')
    
    if query:
        try:
            similar_embeddings = knowledge_base.similarity_search(query)
            similar_embeddings = FAISS.from_documents(documents=similar_embeddings, embedding=OllamaEmbeddings(model="mxbai-embed-large", show_progress=True))
            retriever = similar_embeddings.as_retriever()
            rag_chain = (
                {"context": retriever | format_docs, "question": RunnablePassthrough()}
                | prompt
                | llm
                | StrOutputParser()
            )
            response = rag_chain.invoke(query) + respond_with_sources(query, retriever)
            st.write(response)
        except Exception as e:
            logging.error(f"Error processing query: {e}")
            st.write("An error occurred while processing your query. Please check the logs.")

